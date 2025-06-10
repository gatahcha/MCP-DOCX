# docseditor.py
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocsEditor():
    def __init__(self):
        self._doc = Document()  # Fixed: Use Document() not doc()
        self._title = ""  # Store title for footer use
        # Set default margins
        self.set_margins()
        
    def set_margins(self, top=1.0, bottom=1.0, left=0.75, right=0.75):
        """Set document margins (in inches)
        Default: Top/Bottom: 1", Left/Right: 0.75"
        """
        section = self._doc.sections[0]
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
        
        print(f"Margins set - Top: {top}\", Bottom: {bottom}\", Left: {left}\", Right: {right}\"")
        
    def _format_text(self, run, is_header=False, is_source=False):
        """Apply consistent formatting: black color and Times New Roman
        - Headers: 18pt (changed from 21pt)
        - Sources: 10pt
        - Regular text: 12pt
        """
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        run.font.name = 'Times New Roman'
        if is_header:
            run.font.size = Pt(18)  # Header size (changed from 21pt to 18pt)
        elif is_source:
            run.font.size = Pt(10)  # Source size
        else:
            run.font.size = Pt(12)  # Text size
        
    def add_title(self, title_text, level=0):
        """Add a title to the document - centered, bold, Times New Roman 22pt with full-width black underline"""
        # Store title for footer use
        self._title = title_text
        
        # Create title as a regular paragraph instead of heading to avoid Word's default styles
        title_p = self._doc.add_paragraph()
        title_run = title_p.add_run(title_text)
        
        # Apply custom formatting
        title_run.font.name = 'Times New Roman'
        title_run.font.size = Pt(22)  # Changed from 24pt to 22pt
        title_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        title_run.font.bold = True
        
        # Center the title
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add some spacing after title
        title_p.space_after = Pt(6)
        
        # Add full-width black underline
        self._add_full_underline()
        
        return title_p
    
    def _add_full_underline(self):
        """Add a full-width black underline using border"""
        underline_p = self._doc.add_paragraph()
        
        # Create a paragraph with bottom border (full-width line)
        from docx.oxml.shared import qn
        from docx.oxml import OxmlElement
        
        # Get paragraph properties
        pPr = underline_p._element.get_or_add_pPr()
        
        # Create paragraph borders
        pBdr = OxmlElement('w:pBdr')
        
        # Create bottom border
        bottom_border = OxmlElement('w:bottom')
        bottom_border.set(qn('w:val'), 'single')  # Single line
        bottom_border.set(qn('w:sz'), '18')  # Line thickness (18 = 2.25pt)
        bottom_border.set(qn('w:space'), '1')  # Space from text
        bottom_border.set(qn('w:color'), '000000')  # Black color
        
        pBdr.append(bottom_border)
        pPr.append(pBdr)
        
        # Add small spacing after the line
        underline_p.space_after = Pt(12)
        
        return underline_p
    
    def _add_horizontal_line(self):
        """Add a black horizontal line"""
        p = self._doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a line using underscores or dashes
        run = p.add_run("_" * 50)  # 50 underscores for the line
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black line
        
        return p
    
    def add_paragraph(self, text):
        """Add a paragraph to the document"""
        p = self._doc.add_paragraph()
        run = p.add_run(text)
        self._format_text(run, is_header=False)
        # Justify the paragraph
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p
    
    def add_heading(self, heading_text, level=1):
        """Add a heading to the document - not bold"""
        heading = self._doc.add_heading(heading_text, level)
        # Format the heading
        if heading.runs:
            self._format_text(heading.runs[0], is_header=True)
            heading.runs[0].font.bold = False  # Make heading NOT bold
        # Justify the heading
        heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return heading
    
    def add_source(self, dates, full_source, shorten_source):
        """
        Add source with dates and clickable shortened source
        Format: ("dates", shorten_source) 
        The shorten_source will be clickable to full_source URL
        """
        p = self._doc.add_paragraph()
        
        # Add opening parenthesis
        run1 = p.add_run("(")
        self._format_text(run1, is_source=True)
        
        # Add dates
        run2 = p.add_run(dates)
        self._format_text(run2, is_source=True)
        
        # Add comma and space
        run3 = p.add_run(", ")
        self._format_text(run3, is_source=True)
        
        # Add clickable hyperlink using add_hyperlink method
        hyperlink = self._add_hyperlink_to_paragraph(p, full_source, shorten_source)
        
        # Add closing parenthesis
        run4 = p.add_run(")")
        self._format_text(run4, is_source=True)
        
        # Right-align the source
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        return p
    
    def _add_hyperlink_to_paragraph(self, paragraph, url, text):
        """Add a proper clickable hyperlink to a paragraph"""
        try:
            # Try to create a proper hyperlink
            from docx.oxml.shared import qn
            from docx.oxml import OxmlElement
            
            # Create hyperlink element
            hyperlink_element = OxmlElement('w:hyperlink')
            
            # Create relationship
            part = paragraph.part
            r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
            hyperlink_element.set(qn('r:id'), r_id)
            
            # Create run for the hyperlink
            run_element = OxmlElement('w:r')
            
            # Add run properties for formatting
            run_props = OxmlElement('w:rPr')
            
            # Font settings
            font_element = OxmlElement('w:rFonts')
            font_element.set(qn('w:ascii'), 'Times New Roman')
            run_props.append(font_element)
            
            # Font size (10pt = 20 half-points)
            size_element = OxmlElement('w:sz')
            size_element.set(qn('w:val'), '20')
            run_props.append(size_element)
            
            # Color (black)
            color_element = OxmlElement('w:color')
            color_element.set(qn('w:val'), '000000')
            run_props.append(color_element)
            
            # Underline
            underline_element = OxmlElement('w:u')
            underline_element.set(qn('w:val'), 'single')
            run_props.append(underline_element)
            
            run_element.append(run_props)
            
            # Add text
            text_element = OxmlElement('w:t')
            text_element.text = text
            run_element.append(text_element)
            
            hyperlink_element.append(run_element)
            paragraph._element.append(hyperlink_element)
            
            return hyperlink_element
            
        except Exception as e:
            # Fallback: if hyperlink creation fails, add as underlined text
            print(f"Hyperlink creation failed, using fallback: {e}")
            hyperlink = paragraph.add_run(text)
            self._format_text(hyperlink, is_source=True)
            hyperlink.font.underline = True
            return hyperlink
    
    def add_footer(self):
        """Add footer with title on left and copyright on right"""
        # Access the default section and its footer
        section = self._doc.sections[0]
        footer = section.footer
        
        # Clear any existing footer content
        footer.paragraphs[0].clear()
        
        # Create footer paragraph
        footer_p = footer.paragraphs[0]
        
        left_text = f"{self._title}"
        
        left_run = footer_p.add_run(left_text + "     " + "(c) Generated using MCP_DOCX" )
        left_run.font.name = 'Times New Roman'
        left_run.font.size = Pt(10)
        left_run.font.color.rgb = RGBColor(128, 128, 128)  # Black
        left_run.font.bold = True
        
        # Set paragraph alignment
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return footer_p
    
    def save(self, filename):
        """Save the document"""
        self._doc.save(filename)
        print(f"Document saved as: {filename}")